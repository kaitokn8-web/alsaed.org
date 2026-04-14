import os, re, random
import streamlit as st
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage

load_dotenv()

EMBEDDING_MODEL = "intfloat/multilingual-e5-large"

st.set_page_config(page_title="السَّاعِدُ العِلْمِيُّ", page_icon="🕌", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700&display=swap');
:root{--bg-main:#f4f7f5;--bg-card:#ffffff;--bg-soft:#f1f8f4;--border:#d4e8db;--border-soft:#e8f5e9;--green-dark:#1a3c2e;--green-mid:#2d6a4f;--green-light:#a8d5b5;--gold:#c9a227;--text-main:#1a1a1a;--text-muted:#5a6e62;--text-hint:#8fa897;--shadow:rgba(26,60,46,0.08);--answer-border:#2d6a4f;}
@media(prefers-color-scheme:dark){:root{--bg-main:#0d1a12;--bg-card:#132019;--bg-soft:#1a2e20;--border:#2a4a35;--border-soft:#1e3828;--green-dark:#a8d5b5;--green-mid:#6bbf8a;--green-light:#4a7a5a;--gold:#e8bb3a;--text-main:#e8f0eb;--text-muted:#8ab89a;--text-hint:#5a7a65;--shadow:rgba(0,0,0,0.3);--answer-border:#4a9a6a;}body,.stApp{background-color:var(--bg-main)!important;}}
*{font-family:'Tajawal',sans-serif!important;}
body,.stApp{direction:rtl;background:var(--bg-main);color:var(--text-main);}
.block-container{padding-top:0!important;margin-top:0!important;}
#MainMenu,footer,header{visibility:hidden;}
.header{background:linear-gradient(135deg,var(--green-dark) 0%,var(--green-mid) 100%);padding:22px 40px;border-radius:0 0 24px 24px;margin-bottom:28px;box-shadow:0 4px 24px var(--shadow);text-align:center;}
.header-title{color:var(--gold);font-size:2.2rem;font-weight:700;margin:0;letter-spacing:1px;}
.header-sub{color:var(--green-light);font-size:0.95rem;margin-top:6px;opacity:0.9;}
.hero{background:var(--bg-card);border-radius:16px;padding:28px 36px;text-align:center;margin-bottom:20px;box-shadow:0 2px 16px var(--shadow);border:1px solid var(--border-soft);}
.hero-title{color:var(--green-dark);font-size:1.5rem;font-weight:700;margin-bottom:8px;}
.hero-desc{color:var(--text-muted);font-size:0.98rem;line-height:1.9;max-width:680px;margin:0 auto;}
.stats-row{display:flex;gap:14px;margin-bottom:22px;justify-content:center;}
.stat-card{background:var(--bg-card);border-radius:12px;padding:16px 22px;text-align:center;box-shadow:0 2px 10px var(--shadow);border:1px solid var(--border-soft);flex:1;transition:transform 0.2s;}
.stat-card:hover{transform:translateY(-2px);}
.stat-number{color:var(--green-mid);font-size:1.6rem;font-weight:700;}
.stat-label{color:var(--text-hint);font-size:0.8rem;margin-top:3px;}
.section-title{color:var(--green-dark);font-size:1.05rem;font-weight:700;margin:18px 0 12px;padding-bottom:8px;border-bottom:2px solid var(--border-soft);direction:rtl;text-align:right;}
.answer-box{direction:rtl!important;text-align:right!important;background:var(--bg-card);border-radius:16px;padding:28px 32px;box-shadow:0 2px 14px var(--shadow);border-right:5px solid var(--answer-border);margin-bottom:20px;line-height:2.1;color:var(--text-main)!important;font-size:1.05rem;}
.stButton button{background:var(--bg-soft)!important;color:var(--green-dark)!important;border:1px solid var(--border)!important;border-radius:10px!important;font-family:'Tajawal',sans-serif!important;font-size:0.88rem!important;transition:all 0.2s!important;width:100%!important;padding:8px 12px!important;}
.stButton button:hover{background:var(--green-mid)!important;color:#fff!important;border-color:var(--green-mid)!important;transform:translateY(-1px)!important;box-shadow:0 4px 14px rgba(45,106,79,0.3)!important;}
textarea,input[type="text"]{direction:rtl!important;text-align:right!important;font-family:'Tajawal',sans-serif!important;font-size:1rem!important;border-radius:12px!important;border:2px solid var(--border-soft)!important;background:var(--bg-card)!important;color:var(--text-main)!important;}
[data-testid="stExpander"]{background:var(--bg-soft)!important;border:1px solid var(--border)!important;border-radius:10px!important;margin-bottom:8px!important;}
[data-testid="stExpander"] summary,.streamlit-expanderHeader p{color:var(--green-dark)!important;font-family:'Tajawal',sans-serif!important;font-weight:700!important;font-size:0.95rem!important;}
[data-testid="stExpanderToggleIcon"]{color:var(--green-dark)!important;}
.source-content{direction:rtl;text-align:right;font-family:'Tajawal',sans-serif;font-size:0.93rem;line-height:2;color:var(--text-main);padding:6px 4px;}
.source-book-name{color:var(--green-mid);font-weight:700;font-size:0.95rem;margin-bottom:10px;padding-bottom:6px;border-bottom:1px solid var(--border-soft);}
.source-page-badge{display:inline-block;background:var(--green-mid);color:white;border-radius:20px;padding:2px 10px;font-size:0.78rem;margin-right:8px;vertical-align:middle;}
.stSpinner p{color:var(--green-dark)!important;font-family:'Tajawal',sans-serif!important;}
.stAlert{border-radius:12px!important;direction:rtl!important;}
</style>""", unsafe_allow_html=True)


def build_prompt(context, question):
    return f"""أنت "السَّاعِدُ العِلْمِيُّ"، نظام ذكاء اصطناعي متخصص حصراً في علم العقيدة الإسلامية وفق منهج أهل السنة والجماعة على فهم السلف الصالح.

فيما يلي النصوص المرجعية المستخرجة من كتب أهل السنة — استند إليها في إجابتك:

{context}

---

السؤال: {question}

منهج الإجابة (اتبع هذا الترتيب):
١. تعريف المسألة بإيجاز
٢. الأصل في القرآن والسنة مع ذكر النص إن أمكن
٣. أقوال علماء أهل السنة مقتبسة من النصوص أعلاه، وكل قول يتبعه مباشرة: [اسم الكتاب، ص.رقم]
٤. شرح المسألة وبيان معناها
٥. ذكر الأقوال المخالفة إن وجدت مع موقف أهل السنة منها
٦. خلاصة تلخص مذهب أهل السنة

قواعد إلزامية:
- كل قول تنقله من النصوص أعلاه يعقبه مباشرة: [اسم الكتاب، ص.رقم]
- لا تذكر قولاً لعالم إلا إذا كان موجوداً في النصوص المرفقة
- لا تخترع مصادر أو صفحات
- اكتب بالعربية الفصحى بأسلوب علمي رصين
"""


BOOK_NAMES = {
    "tahawiyya321.docx":                    "العقيدة الطحاوية — الطحاوي (ت.321)",
    "tahawiyya792_1.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.1",
    "tahawiyya792_2.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.2",
    "tahawiyya792_3.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.3",
    "jawab_sahih_728_1.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.1",
    "jawab_sahih_728_2.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.2",
    "jawab_sahih_728_3.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.3",
    "jawab_sahih_728_4.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.4",
    "jawab_sahih_728_5.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.5",
    "jawab_sahih_728_6.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.6",
    "jawab_sahih_728_7.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.7",
    "itiqad_lalkai_418_1.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.1",
    "itiqad_lalkai_418_2.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.2",
    "itiqad_lalkai_418_3.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.3",
    "itiqad_lalkai_418_4.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.4",
    "itiqad_lalkai_418_5.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.5",
    "itiqad_lalkai_418_6.docx":             "كرامات الأولياء — اللالكائي (ت.418) ج.6",
    "dar_taarez_728_1.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.1",
    "dar_taarez_728_2.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.2",
    "dar_taarez_728_3.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.3",
    "dar_taarez_728_4.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.4",
    "dar_taarez_728_5.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.5",
    "dar_taarez_728_6.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.6",
    "dar_taarez_728_7.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.7",
    "dar_taarez_728_8.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.8",
    "radd_shadhili_728.docx":               "الرد على الشاذلي في حزبيه — ابن تيمية (ت.728)",
    "sharh_asfahaniyya_728.docx":           "شرح العقيدة الأصفهانية — ابن تيمية (ت.728)",
    "ubudiyya_728.docx":                    "العبودية — ابن تيمية (ت.728)",
    "nubuwwat_728_1.docx":                  "النبوات — ابن تيمية (ت.728) ج.1",
    "nubuwwat_728_2.docx":                  "النبوات — ابن تيمية (ت.728) ج.2",
    "nubuwwat_728_3.docx":                  "النبوات — ابن تيمية (ت.728) ج.3",
    "nubuwwat_728_4.docx":                  "النبوات — ابن تيمية (ت.728) ج.4",
    "dhimma_751.docx":                      "أحكام أهل الذمة — ابن القيم (ت.751)",
    "asma_sifat_1421.docx":                 "أسماء الله وصفاته وموقف أهل السنة منها — ابن عثيمين (ت.1421)",
    "ishtiqaq_asma_337.docx":               "اشتقاق أسماء الله — الزجاج (ت.337)",
    "iqtidaa_sirat_728.docx":               "اقتضاء الصراط المستقيم مخالفة أصحاب الجحيم — ابن تيمية (ت.728)",
    "ibanah_ashari_324.docx":               "الإبانة عن أصول الديانة — الأشعري (ت.324)",
    "athar_marwiyya_tamimi.docx":           "الآثار المروية في صفة المعية — محمد التميمي",
    "itiqad_ibn_abi_yala_526.docx":         "الاعتقاد — ابن أبي يعلى (ت.526)",
    "iqtisad_itiqad_600.docx":              "الاقتصاد في الاعتقاد — عبد الغني المقدسي (ت.600)",
    "iman_awsat_728.docx":                  "الإيمان الأوسط — ابن تيمية (ت.728)",
    "kitab_iman_224.docx":                  "كتاب الإيمان — القاسم بن سلام (ت.224)",
    "tanbih_radd_377.docx":                 "التنبيه والرد على أهل الأهواء والبدع — العسقلاني (ت.377)",
    "hayda_kinani_240.docx":                "الحيدة والاعتذار في الرد على من قال بخلق القرآن — الكناني (ت.240)",
    "radd_jahmiyya_241.docx":               "الرد على الجهمية والزنادقة — أحمد بن حنبل (ت.241)",
    "radd_mantiqiyyin_728.docx":            "الرد على المنطقيين — ابن تيمية (ت.728)",
    "risala_akmaliyya_728.docx":            "الرسالة الأكملية في ما يجب لله من صفات الكمال — ابن تيمية (ت.728)",
    "risala_tadmuriyya_728.docx":           "الرسالة التدمرية — ابن تيمية (ت.728)",
    "ruh_751.docx":                         "الروح — ابن القيم (ت.751)",
    "sarim_maslul_728.docx":                "الصارم المسلول على شاتم الرسول — ابن تيمية (ت.728)",
    "sawaiq_mursala_751.docx":              "الصواعق المرسلة في الرد على الجهمية والمعطلة — ابن القيم (ت.751)",
    "aqida_khallal_241.docx":               "العقيدة رواية أبي بكر الخلال — أحمد بن حنبل (ت.241)",
    "awasim_qawasim_543.docx":              "العواصم من القواصم — ابن العربي (ت.543)",
    "fatwa_hamawiyya_728.docx":             "الفتوى الحموية الكبرى — ابن تيمية (ت.728)",
    "farq_bayn_firaq_429.docx":             "الفرق بين الفرق وبيان الفرقة الناجية — الإسفراييني (ت.429)",
    "furqan_awliya_728.docx":               "الفرقان بين أولياء الرحمن وأولياء الشيطان — ابن تيمية (ت.728)",
    "tahrim_nazar_620.docx":                "تحريم النظر في كتب الكلام — ابن قدامة المقدسي (ت.620)",
    "tafsir_asma_husna_311.docx":           "تفسير أسماء الله الحسنى — الزجاج (ت.311)",
    "itiqad_salaf_huruf_676.docx":          "جزء فيه ذكر اعتقاد السلف في الحروف والأصوات — النووي (ت.676)",
    "dhamm_kalam_481.docx":                 "ذم الكلام وأهله — الهروي (ت.481)",
    "risala_sijzi_444.docx":                "رسالة السجزي إلى أهل زبيد في الرد على من أنكر الحرف والصوت — السجزي (ت.444)",
    "risala_thaghr_ashari_324.docx":        "رسالة إلى أهل الثغر بباب الأبواب — الأشعري (ت.324)",
    "risala_quran_ghayr_makhlooq_285.docx": "رسالة في أن القرآن غير مخلوق — أبو إسحاق البغدادي (ت.285)",
    "sharh_wasitiyya_harras_1395.docx":     "شرح العقيدة الواسطية — هرّاس (ت.1395)",
    "shifa_alil_751.docx":                  "شفاء العليل في مسائل القضاء والقدر والحكمة والتعليل — ابن القيم (ت.751)",
    "fadaih_batiniyya_505.docx":            "فضائح الباطنية — الغزالي (ت.505)",
    "qawl_falasifa_yoonan_khalaf.docx":     "قول الفلاسفة اليونان الوثنيين في توحيد الربوبية — سعود الخلف",
    "lamaat_itiqad_620.docx":               "لمعة الاعتقاد — ابن قدامة المقدسي (ت.620)",
    "mustalah_aqaid_hamad.docx":            "مصطلحات في كتب العقائد — محمد الحمد",
    "maarij_qubul_1377.docx":               "معارج القبول بشرح سلم الوصول — الحكمي (ت.1377)",
    "mujam_manahi_1429.docx":               "معجم المناهي اللفظية — بكر أبو زيد (ت.1429)",
    "maqala_tateel_jad_tamimi.docx":        "مقالة التعطيل والجعد بن درهم — محمد التميمي",
    "mawqif_ibn_taymiyya_ashaira_mahmoud.docx": "موقف ابن تيمية من الأشاعرة — عبد الرحمن المحمود",
    "qasidat_nuniyya_751.docx":             "متن القصيدة النونية — ابن القيم (ت.751)",
    "nuniyya_qahtani_378.docx":             "القصيدة النونية — القحطاني (ت.378)",
    "hidayat_hayara_751.docx":              "هداية الحيارى في أجوبة اليهود والنصارى — ابن القيم (ت.751)",
}


@st.cache_resource(show_spinner=False)
def load_db():
    import shutil, re as _re
    from docx import Document as DocxDocument
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LCDocument

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        encode_kwargs={"normalize_embeddings": True}
    )

    needs_build = (
        not os.path.exists("vectorstore")
        or not os.path.exists(os.path.join("vectorstore", "chroma.sqlite3"))
    )

    if needs_build and os.path.exists("books"):
        PATTERNS = [
            _re.compile(r'\(([\d\u0660-\u0669]+)/[\d\u0660-\u0669]+\)'),
            _re.compile(r'\[ص[:\s]*([\d\u0660-\u0669]+)\]'),
            _re.compile(r'ص\.([\d\u0660-\u0669]+)'),
        ]

        def read_docx(path, book_name):
            doc = DocxDocument(path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            current_page = "1"
            result = []
            for i in range(0, len(paragraphs), 15):
                text = "\n".join(paragraphs[i:i+15])
                for pat in PATTERNS:
                    m = pat.search(text)
                    if m:
                        current_page = m.group(1)
                        break
                if len(text) >= 30:
                    result.append(LCDocument(
                        page_content=text,
                        metadata={"book": book_name, "page": current_page,
                                  "source": f"{book_name} ص.{current_page}"}
                    ))
            return result

        all_docs = []
        for filename, name in BOOK_NAMES.items():
            path = os.path.join("books", filename)
            if os.path.exists(path):
                try:
                    all_docs.extend(read_docx(path, name))
                except Exception:
                    pass

        if all_docs:
            chunks = RecursiveCharacterTextSplitter(
                chunk_size=600, chunk_overlap=80,
                separators=["\n\n", "\n", ".", " "]
            ).split_documents(all_docs)
            if os.path.exists("vectorstore"):
                shutil.rmtree("vectorstore")
            Chroma.from_documents(documents=chunks, embedding=embeddings, persist_directory="vectorstore")

    db = Chroma(persist_directory="vectorstore", embedding_function=embeddings)
    return db, embeddings


@st.cache_resource(show_spinner=False)
def load_llm():
    return ChatAnthropic(
        model="claude-sonnet-4-20250514",
        api_key=os.getenv("ANTHROPIC_API_KEY"),
        max_tokens=2500
    )


def search_all_books(question, db, top_books=15):
    """
    يبحث في كل كتاب على حدة — يضمن تمثيل كل كتاب بالتساوي.
    """
    all_data = db.get(include=["metadatas"])
    books = list({m["book"] for m in all_data["metadatas"] if "book" in m})

    results_with_scores = []
    for book in books:
        try:
            docs = db.similarity_search_with_score(
                question, k=1,
                filter={"book": book}
            )
            if docs:
                results_with_scores.append(docs[0])
        except Exception:
            pass

    results_with_scores.sort(key=lambda x: x[1])
    return [doc for doc, score in results_with_scores[:top_books]]


def ask(question, db, llm):
    docs = search_all_books(question, db, top_books=15)
    if not docs:
        return "لم يُعثر على نصوص ذات صلة في قاعدة البيانات.", []

    context = "\n\n".join([
        f"[{d.metadata.get('source','؟')}]\n{d.page_content}" for d in docs
    ])
    response = llm.invoke([HumanMessage(content=build_prompt(context, question))])
    answer = response.content if hasattr(response, "content") else str(response)
    return answer, docs


st.markdown("""
<div class="header">
    <div class="header-title">🕌 السَّاعِدُ العِلْمِيُّ</div>
    <div class="header-sub">مساعد علمي متخصص في عقيدة أهل السنة والجماعة</div>
</div>""", unsafe_allow_html=True)

st.markdown("""
<div class="hero">
    <div class="hero-title">اسأل عن أي مسألة علمية</div>
    <div class="hero-desc">مساعد علمي يبحث في كتب أهل السنة ويجيبك إجابة موثقة</div>
</div>""", unsafe_allow_html=True)

db, embeddings = load_db()
llm = load_llm()

try:
    books_count  = len([f for f in os.listdir("books") if f.endswith((".pdf",".docx",".doc"))]) if os.path.exists("books") else 76
    chunks_count = len(db.get()['ids']) if db else 80000
except Exception:
    books_count, chunks_count = 76, 80000

def to_arabic(n):
    return str(n).translate(str.maketrans('0123456789','٠١٢٣٤٥٦٧٨٩'))

st.markdown(f"""
<div class="stats-row">
<div class="stat-card"><div class="stat-number">{to_arabic(books_count)}</div><div class="stat-label">كتاب مفهرس</div></div>
<div class="stat-card"><div class="stat-number">+{to_arabic(chunks_count)}</div><div class="stat-label">مقطع نصي</div></div>
<div class="stat-card"><div class="stat-number">١٠٠٪</div><div class="stat-label">من المصادر الأصلية</div></div>
</div>""", unsafe_allow_html=True)

if db is None:
    st.error("⚠️ قاعدة البيانات غير موجودة — شغّل 1_ingest.py أولاً")
    st.stop()

ALL_EXAMPLES = [
    "ما معنى توحيد الألوهية؟","ما موقف أهل السنة من صفة الاستواء؟",
    "ما الدليل على إثبات علو الله؟","ما الفرق بين المعتزلة وأهل السنة في الصفات؟",
    "ما اعتقاد أهل السنة في القرآن الكريم؟","ما موقف أهل السنة من الصحابة الكرام؟",
    "ما معنى توحيد الربوبية؟","ما موقف أهل السنة من القدر؟",
    "ما حكم من أنكر صفات الله؟","ما الفرق بين الأشاعرة وأهل السنة؟",
    "ما موقف أهل السنة من الإيمان؟","ما معنى الولاء والبراء؟",
    "ما دلالة آية الكرسي على صفات الله؟","ما معنى الأسماء الحسنى وما حكم الإلحاد فيها؟",
]

if "examples_shown" not in st.session_state:
    st.session_state.examples_shown = random.sample(ALL_EXAMPLES, 3)

st.markdown('<div class="section-title">💡 أسئلة مقترحة</div>', unsafe_allow_html=True)
cols = st.columns(3)
for i, ex in enumerate(st.session_state.examples_shown):
    if cols[i].button(ex, key=f"ex_{i}"):
        st.session_state.q = ex
        st.rerun()

st.markdown('<div class="section-title">🔍 اكتب سؤالك</div>', unsafe_allow_html=True)
question = st.text_input(label="", value=st.session_state.get("q",""),
    placeholder="اكتب سؤالك العقدي هنا... مثال: ما معنى توحيد الربوبية؟", key="question_input")

col_l, col_m, col_r = st.columns([1,2,1])
with col_m:
    search = st.button("🔍 ابحث في كتب العقيدة", type="primary", use_container_width=True)

trigger = search or (question.strip() and question != st.session_state.get("last_q",""))

if trigger:
    if not question.strip():
        st.warning("⚠️ الرجاء كتابة سؤال أولاً")
    else:
        st.session_state.last_q = question
        st.session_state.pop("q", None)

        with st.spinner("⏳ جاري البحث في كتب أهل السنة..."):
            answer, docs = ask(question, db, llm)

        st.markdown('<div class="section-title">📖 الإجابة</div>', unsafe_allow_html=True)

        answer_html = re.sub(r'### (.+)', r'<h3 style="color:var(--green-dark);">\1</h3>', answer)
        answer_html = re.sub(r'## (.+)', r'<h2 style="color:var(--green-dark);font-size:1.2rem;">\1</h2>', answer_html)
        answer_html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', answer_html)
        answer_html = re.sub(r'\[([^\]]+؟?,?\s*ص\.[\d\u0660-\u0669]+)\]',
            r'<span style="color:var(--green-mid);font-weight:700;font-size:0.88em;">[\1]</span>', answer_html)
        answer_html = answer_html.replace('\n','<br>')

        st.markdown(f'<div class="answer-box">{answer_html}</div>', unsafe_allow_html=True)

        with st.expander("📋 انسخ نص الإجابة كاملاً"):
            st.code(answer, language=None)

        st.markdown('<div class="section-title">📚 النصوص المصدرية المستخدمة في الإجابة</div>', unsafe_allow_html=True)

        seen, unique_docs = set(), []
        for doc in docs:
            key = doc.metadata.get("source","")
            if key not in seen:
                seen.add(key)
                unique_docs.append(doc)

        if unique_docs:
            for doc in unique_docs:
                book   = doc.metadata.get("book",  "كتاب غير معروف")
                page   = doc.metadata.get("page",  "—")
                source = doc.metadata.get("source", book)
                with st.expander(f"📖  {source}"):
                    st.markdown(
                        f'<div class="source-content"><div class="source-book-name">{book}'
                        f'<span class="source-page-badge">ص. {page}</span></div>'
                        f'{doc.page_content}</div>', unsafe_allow_html=True)
        else:
            st.info("لم يُعثر على نصوص مصدرية مرتبطة بهذا السؤال في قاعدة البيانات.")

st.markdown("""
<div style="text-align:center;padding:16px 0 4px;direction:rtl;">
<span style="color:var(--text-hint);font-size:0.85rem;">للملاحظات والاستفسارات: </span>
<a href="https://t.me/steripro" target="_blank"
   style="background:#229ED9;color:white;text-decoration:none;padding:5px 16px;border-radius:50px;
          margin:4px;display:inline-block;font-family:Tajawal,sans-serif;font-size:0.85rem;font-weight:700;">✈️ تيليجرام</a>
</div>
<div style="text-align:center;padding:6px 0 20px;">
<span style="color:var(--text-hint);font-size:0.75rem;opacity:0.6;">نسخة تجريبية — السَّاعِدُ العِلْمِيُّ</span>
</div>""", unsafe_allow_html=True)