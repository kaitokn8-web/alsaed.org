import os
import re
import shutil
from dotenv import load_dotenv
load_dotenv()

from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

BOOK_NAMES = {
    "tahawiyya321.docx":                    "العقيدة الطحاوية — الطحاوي (ت.321)",
    "tahawiyya792_1.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.1",
    "tahawiyya792_2.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.2",
    "tahawiyya792_3.docx":                  "شرح الطحاوية — ابن أبي العز (ت.792) ج.3",
    "jawab_sahih_728_1.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.1",
    "jawab_sahih_728_2.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.2",
    "jawab_sahih_728_3.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.3",
    "jawab_sahih_728_4.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.4",
    "jawab_sahih_728_5.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.5",
    "jawab_sahih_728_6.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.6",
    "jawab_sahih_728_7.docx":               "الجواب الصحيح — ابن تيمية (ت.728) ج.7",
    "itiqad_lalkai_418_1.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.1",
    "itiqad_lalkai_418_2.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.2",
    "itiqad_lalkai_418_3.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.3",
    "itiqad_lalkai_418_4.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.4",
    "itiqad_lalkai_418_5.docx":             "شرح أصول اعتقاد أهل السنة — اللالكائي (ت.418) ج.5",
    "itiqad_lalkai_418_6.docx":             "كرامات الأولياء — اللالكائي (ت.418) ج.6",
    "dar_taarez_728_1.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.1",
    "dar_taarez_728_2.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.2",
    "dar_taarez_728_3.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.3",
    "dar_taarez_728_4.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.4",
    "dar_taarez_728_5.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.5",
    "dar_taarez_728_6.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.6",
    "dar_taarez_728_7.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.7",
    "dar_taarez_728_8.docx":                "درء تعارض العقل والنقل — ابن تيمية (ت.728) ج.8",
    "radd_shadhili_728.docx":               "الرد على الشاذلي في حزبيه — ابن تيمية (ت.728)",
    "sharh_asfahaniyya_728.docx":           "شرح العقيدة الأصفهانية — ابن تيمية (ت.728)",
    "ubudiyya_728.docx":                    "العبودية — ابن تيمية (ت.728)",
    "nubuwwat_728_1.docx":                  "النبوات — ابن تيمية (ت.728) ج.1",
    "nubuwwat_728_2.docx":                  "النبوات — ابن تيمية (ت.728) ج.2",
    "nubuwwat_728_3.docx":                  "النبوات — ابن تيمية (ت.728) ج.3",
    "nubuwwat_728_4.docx":                  "النبوات — ابن تيمية (ت.728) ج.4",
    "dhimma_751.docx":                      "أحكام أهل الذمة — ابن القيم (ت.751)",
    "asma_sifat_1421.docx":                 "أسماء الله وصفاته وموقف أهل السنة منها — ابن عثيمين (ت.1421)",
    "ishtiqaq_asma_337.docx":               "اشتقاق أسماء الله — الزجاج (ت.337)",
    "iqtidaa_sirat_728.docx":               "اقتضاء الصراط المستقيم مخالفة أصحاب الجحيم — ابن تيمية (ت.728)",
    "ibanah_ashari_324.docx":               "الإبانة عن أصول الديانة — الأشعري (ت.324)",
    "athar_marwiyya_tamimi.docx":           "الآثار المروية في صفة المعية — محمد التميمي",
    "itiqad_ibn_abi_yala_526.docx":         "الاعتقاد — ابن أبي يعلى (ت.526)",
    "iqtisad_itiqad_600.docx":              "الاقتصاد في الاعتقاد — عبد الغني المقدسي (ت.600)",
    "iman_awsat_728.docx":                  "الإيمان الأوسط — ابن تيمية (ت.728)",
    "kitab_iman_224.docx":                  "كتاب الإيمان — القاسم بن سلام (ت.224)",
    "tanbih_radd_377.docx":                 "التنبيه والرد على أهل الأهواء والبدع — العسقلاني (ت.377)",
    "hayda_kinani_240.docx":                "الحيدة والاعتذار في الرد على من قال بخلق القرآن — الكناني (ت.240)",
    "radd_jahmiyya_241.docx":               "الرد على الجهمية والزنادقة — أحمد بن حنبل (ت.241)",
    "radd_mantiqiyyin_728.docx":            "الرد على المنطقيين — ابن تيمية (ت.728)",
    "risala_akmaliyya_728.docx":            "الرسالة الأكملية في ما يجب لله من صفات الكمال — ابن تيمية (ت.728)",
    "risala_tadmuriyya_728.docx":           "الرسالة التدمرية — ابن تيمية (ت.728)",
    "ruh_751.docx":                         "الروح — ابن القيم (ت.751)",
    "sarim_maslul_728.docx":                "الصارم المسلول على شاتم الرسول — ابن تيمية (ت.728)",
    "sawaiq_mursala_751.docx":              "الصواعق المرسلة في الرد على الجهمية والمعطلة — ابن القيم (ت.751)",
    "aqida_khallal_241.docx":               "العقيدة رواية أبي بكر الخلال — أحمد بن حنبل (ت.241)",
    "awasim_qawasim_543.docx":              "العواصم من القواصم — ابن العربي (ت.543)",
    "fatwa_hamawiyya_728.docx":             "الفتوى الحموية الكبرى — ابن تيمية (ت.728)",
    "farq_bayn_firaq_429.docx":             "الفرق بين الفرق وبيان الفرقة الناجية — الإسفراييني (ت.429)",
    "furqan_awliya_728.docx":               "الفرقان بين أولياء الرحمن وأولياء الشيطان — ابن تيمية (ت.728)",
    "tahrim_nazar_620.docx":                "تحريم النظر في كتب الكلام — ابن قدامة المقدسي (ت.620)",
    "tafsir_asma_husna_311.docx":           "تفسير أسماء الله الحسنى — الزجاج (ت.311)",
    "itiqad_salaf_huruf_676.docx":          "جزء فيه ذكر اعتقاد السلف في الحروف والأصوات — النووي (ت.676)",
    "dhamm_kalam_481.docx":                 "ذم الكلام وأهله — الهروي (ت.481)",
    "risala_sijzi_444.docx":                "رسالة السجزي إلى أهل زبيد في الرد على من أنكر الحرف والصوت — السجزي (ت.444)",
    "risala_thaghr_ashari_324.docx":        "رسالة إلى أهل الثغر بباب الأبواب — الأشعري (ت.324)",
    "risala_quran_ghayr_makhlooq_285.docx": "رسالة في أن القرآن غير مخلوق — أبو إسحاق البغدادي (ت.285)",
    "sharh_wasitiyya_harras_1395.docx":     "شرح العقيدة الواسطية — هرّاس (ت.1395)",
    "shifa_alil_751.docx":                  "شفاء العليل في مسائل القضاء والقدر والحكمة والتعليل — ابن القيم (ت.751)",
    "fadaih_batiniyya_505.docx":            "فضائح الباطنية — الغزالي (ت.505)",
    "qawl_falasifa_yoonan_khalaf.docx":     "قول الفلاسفة اليونان الوثنيين في توحيد الربوبية — سعود الخلف",
    "lamaat_itiqad_620.docx":               "لمعة الاعتقاد — ابن قدامة المقدسي (ت.620)",
    "mustalah_aqaid_hamad.docx":            "مصطلحات في كتب العقائد — محمد الحمد",
    "maarij_qubul_1377.docx":               "معارج القبول بشرح سلم الوصول — الحكمي (ت.1377)",
    "mujam_manahi_1429.docx":               "معجم المناهي اللفظية — بكر أبو زيد (ت.1429)",
    "maqala_tateel_jad_tamimi.docx":        "مقالة التعطيل والجعد بن درهم — محمد التميمي",
    "mawqif_ibn_taymiyya_ashaira_mahmoud.docx": "موقف ابن تيمية من الأشاعرة — عبد الرحمن المحمود",
    "qasidat_nuniyya_751.docx":             "متن القصيدة النونية — ابن القيم (ت.751)",
    "nuniyya_qahtani_378.docx":             "القصيدة النونية — القحطاني (ت.378)",
    "hidayat_hayara_751.docx":              "هداية الحيارى في أجوبة اليهود والنصارى — ابن القيم (ت.751)",
}

# ── الكتب الصغيرة التي تحتاج تضخيماً (أقل من 50 مقطعاً) ──
# نضاعف مقاطعها 4 مرات لتظهر في نتائج البحث بالتساوي مع الكتب الكبيرة
BOOST_BOOKS = {
    "mustalah_aqaid_hamad.docx",        # 37 مقطع
    "athar_marwiyya_tamimi.docx",       # 9 مقاطع
    "maqala_tateel_jad_tamimi.docx",    # 24 مقطع
    "qawl_falasifa_yoonan_khalaf.docx", # 5 مقاطع
    "risala_akmaliyya_728.docx",        # 10 مقاطع
    "risala_quran_ghayr_makhlooq_285.docx", # 4 مقاطع
    "tahawiyya321.docx",                # 4 مقاطع
    "aqida_khallal_241.docx",           # 4 مقاطع
    "tafsir_asma_husna_311.docx",       # 6 مقاطع
    "lamaat_itiqad_620.docx",           # 6 مقاطع
    "tahrim_nazar_620.docx",            # 6 مقاطع
    "asma_sifat_1421.docx",             # 7 مقاطع
    "itiqad_ibn_abi_yala_526.docx",     # 7 مقاطع
    "nuniyya_qahtani_378.docx",         # 6 مقاطع
    "risala_thaghr_ashari_324.docx",    # 26 مقطع
    "itiqad_salaf_huruf_676.docx",      # 12 مقطع
    "hayda_kinani_240.docx",            # 12 مقطع
    "kitab_iman_224.docx",              # 13 مقطع
    "fatwa_hamawiyya_728.docx",         # 18 مقطع
    "risala_tadmuriyya_728.docx",       # 19 مقطع
    "ibanah_ashari_324.docx",           # 33 مقطع
    "fadaih_batiniyya_505.docx",        # 31 مقطع
    "iqtisad_itiqad_600.docx",          # 32 مقطع
    "tanbih_radd_377.docx",             # 25 مقطع
    "radd_jahmiyya_241.docx",           # 26 مقطع
}

BOOST_FACTOR = 4  # عدد مرات التضخيم للكتب الصغيرة


def read_docx(path, book_name):
    PATTERNS = [
        re.compile(r'\(([\d\u0660-\u0669]+)/([\d\u0660-\u0669]+)\)'),
        re.compile(r'\[ص[:\s]*([\d\u0660-\u0669]+)\]'),
        re.compile(r'ص\.([\d\u0660-\u0669]+)'),
    ]

    doc          = DocxDocument(path)
    paragraphs   = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    current_page = "1"
    page_counter = 1
    result       = []

    PARA_PER_CHUNK = 15
    for i in range(0, len(paragraphs), PARA_PER_CHUNK):
        group = paragraphs[i : i + PARA_PER_CHUNK]
        text  = "\n".join(group)

        found = False
        for pat in PATTERNS:
            m = pat.search(text)
            if m:
                current_page = m.group(1)
                found = True
                break

        if not found:
            page_counter += 1
            if current_page == "1":
                current_page = str(page_counter)

        if len(text) < 30:
            continue

        result.append(Document(
            page_content=text,
            metadata={
                "book":   book_name,
                "page":   current_page,
                "source": f"{book_name} ص.{current_page}"
            }
        ))

    print(f"  ✓ {book_name}: {len(result)} مقطع")
    return result


def build():
    print("=" * 60)
    print("📚 السَّاعِدُ العِلْمِيُّ — بناء قاعدة البيانات")
    print("=" * 60)

    print("\n📖 جاري قراءة الكتب...\n")
    all_docs = []

    for filename, name in BOOK_NAMES.items():
        path = os.path.join("books", filename)
        if not os.path.exists(path):
            print(f"  ⚠  غير موجود: {filename}")
            continue
        try:
            docs = read_docx(path, name)
            # تضخيم الكتب الصغيرة
            if filename in BOOST_BOOKS:
                docs = docs * BOOST_FACTOR
                print(f"  🔼 تضخيم {name}: {len(docs)} مقطع بعد التضخيم")
            all_docs.extend(docs)
        except Exception as e:
            print(f"  ✗  خطأ في {filename}: {e}")

    if not all_docs:
        print("\n❌ لا توجد كتب في مجلد books/")
        return

    print(f"\n📄 إجمالي الأجزاء المقروءة: {len(all_docs)}")

    print("\n✂  جاري تقطيع النصوص...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=80,
        separators=["\n\n", "\n", ".", " "]
    )
    chunks = splitter.split_documents(all_docs)
    print(f"✂  تم التقطيع إلى {len(chunks)} مقطع")

    print("\n⏳ جاري بناء قاعدة البيانات المتجهة...")
    print("   (هذا قد يأخذ عدة دقائق أول مرة)")

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )

    if os.path.exists("vectorstore"):
        shutil.rmtree("vectorstore")

    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory="vectorstore"
    )

    print(f"\n{'=' * 60}")
    print(f"✅  اكتمل البناء!")
    print(f"   الكتب المفهرسة : {len(BOOK_NAMES)}")
    print(f"   المقاطع النصية : {len(chunks)}")
    print(f"{'=' * 60}\n")


if __name__ == "__main__":
    build()